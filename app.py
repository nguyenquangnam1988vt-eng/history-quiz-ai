import streamlit as st
import json
import sqlite3
import random
import string
import re
from datetime import datetime
import io
import docx
import PyPDF2
import google.generativeai as genai
import os
import pandas as pd

# ==================== CẤU HÌNH ====================
st.set_page_config(
    page_title="Quiz Lịch Sử AI",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS tùy chỉnh
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 2rem;
        padding: 20px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
    }
    .quiz-card {
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        border-left: 5px solid #3B82F6;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .score-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        margin: 1rem 0;
    }
    .student-info-card {
        background-color: #e3f2fd;
        padding: 20px;
        border-radius: 10px;
        margin: 10px 0;
        border-left: 5px solid #2196F3;
    }
    .warning-card {
        background-color: #fff3cd;
        padding: 20px;
        border-radius: 10px;
        margin: 10px 0;
        border-left: 5px solid #ffc107;
    }
    .stButton > button {
        width: 100%;
        background-color: #3B82F6;
        color: white;
        font-weight: bold;
        border: none;
        padding: 12px;
        border-radius: 8px;
        font-size: 1.1em;
    }
    .stButton > button:hover {
        background-color: #2563EB;
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    .required-field::after {
        content: " *";
        color: red;
    }
    .rank-1 { background-color: #FFD700 !important; color: black; }
    .rank-2 { background-color: #C0C0C0 !important; color: black; }
    .rank-3 { background-color: #CD7F32 !important; color: white; }
    .answer-correct { background-color: #d4edda !important; border-left: 5px solid #28a745 !important; }
    .answer-wrong { background-color: #f8d7da !important; border-left: 5px solid #dc3545 !important; }
    .badge {
        display: inline-block;
        padding: 5px 10px;
        border-radius: 20px;
        font-size: 0.9em;
        font-weight: bold;
        margin: 2px;
    }
    .badge-success { background-color: #198754; color: white; }
    .badge-warning { background-color: #ffc107; color: black; }
    .badge-danger { background-color: #dc3545; color: white; }
    .badge-info { background-color: #0dcaf0; color: white; }
</style>
""", unsafe_allow_html=True)

# ==================== ĐĂNG NHẬP ADMIN ====================
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

def login_admin():
    """Hiển thị form đăng nhập admin trên sidebar"""
    st.sidebar.markdown("---")
    if not st.session_state.logged_in:
        with st.sidebar:
            st.subheader("👨‍🏫 ĐĂNG NHẬP GIÁO VIÊN")
            admin_user = st.text_input("Tài khoản", key="admin_user")
            admin_pass = st.text_input("Mật khẩu", type="password", key="admin_pass")
            if st.button("Đăng nhập", use_container_width=True):
                if admin_user == "admin" and admin_pass == "duyen123":
                    st.session_state.logged_in = True
                    st.rerun()
                else:
                    st.error("Sai tài khoản hoặc mật khẩu!")
    else:
        st.sidebar.success("✅ Đã đăng nhập với quyền Giáo viên")
        if st.sidebar.button("Đăng xuất", use_container_width=True):
            st.session_state.logged_in = False
            st.rerun()

# ==================== DATABASE MIGRATION ====================
def migrate_database():
    """Cập nhật cấu trúc database khi có thay đổi"""
    conn = sqlite3.connect('quiz_system.db')
    c = conn.cursor()
    
    try:
        # Kiểm tra xem bảng results đã tồn tại chưa
        c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='results'")
        if not c.fetchone():
            # Tạo bảng mới với đầy đủ cột
            c.execute('''CREATE TABLE results
                         (id INTEGER PRIMARY KEY AUTOINCREMENT,
                          quiz_code TEXT,
                          student_name TEXT NOT NULL,
                          class_name TEXT NOT NULL,
                          student_id TEXT DEFAULT '',
                          score INTEGER,
                          total_questions INTEGER,
                          percentage REAL DEFAULT 0,
                          grade TEXT DEFAULT '',
                          submitted_at TIMESTAMP)''')
            print("✅ Tạo bảng results mới")
        else:
            # Kiểm tra và thêm cột nếu thiếu
            c.execute("PRAGMA table_info(results)")
            columns = [col[1] for col in c.fetchall()]
            
            columns_to_add = [
                ('class_name', 'TEXT NOT NULL DEFAULT ""'),
                ('student_id', 'TEXT DEFAULT ""'),
                ('percentage', 'REAL DEFAULT 0'),
                ('grade', 'TEXT DEFAULT ""')
            ]
            
            for col_name, col_type in columns_to_add:
                if col_name not in columns:
                    print(f"🔄 Thêm cột {col_name}...")
                    c.execute(f"ALTER TABLE results ADD COLUMN {col_name} {col_type}")
            
            print("✅ Database migration hoàn tất!")
        
    except Exception as e:
        print(f"⚠️ Lỗi migration: {e}")
        # Nếu lỗi nặng, tạo lại bảng
        try:
            c.execute('DROP TABLE IF EXISTS results')
            c.execute('''CREATE TABLE results
                         (id INTEGER PRIMARY KEY AUTOINCREMENT,
                          quiz_code TEXT,
                          student_name TEXT NOT NULL,
                          class_name TEXT NOT NULL,
                          student_id TEXT DEFAULT '',
                          score INTEGER,
                          total_questions INTEGER,
                          percentage REAL DEFAULT 0,
                          grade TEXT DEFAULT '',
                          submitted_at TIMESTAMP)''')
            print("🔄 Tạo lại bảng results...")
        except Exception as e2:
            print(f"❌ Lỗi nặng: {e2}")
    
    conn.commit()
    conn.close()

def check_and_fix_database():
    """Tự động kiểm tra và sửa lỗi database"""
    try:
        conn = sqlite3.connect('quiz_system.db')
        c = conn.cursor()
        
        # Kiểm tra bảng quizzes
        c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='quizzes'")
        if not c.fetchone():
            print("⚠️ Bảng quizzes không tồn tại, đang tạo...")
            reset_quizzes_table()
            conn.close()
            return
        
        c.execute("PRAGMA table_info(quizzes)")
        columns = [col[1] for col in c.fetchall()]
        
        required_columns = [
            'id', 'quiz_code', 'title', 'subject', 'created_at', 
            'question_count', 'is_active', 'difficulty'
        ]
        
        missing_columns = [col for col in required_columns if col not in columns]
        
        if missing_columns:
            print(f"⚠️ Thiếu cột trong quizzes: {missing_columns}")
            
            # Thêm các cột bị thiếu
            for col in missing_columns:
                try:
                    if col == 'difficulty':
                        c.execute(f"ALTER TABLE quizzes ADD COLUMN {col} TEXT DEFAULT 'medium'")
                    elif col == 'is_active':
                        c.execute(f"ALTER TABLE quizzes ADD COLUMN {col} BOOLEAN DEFAULT 1")
                    elif col == 'question_count':
                        c.execute(f"ALTER TABLE quizzes ADD COLUMN {col} INTEGER DEFAULT 0")
                    else:
                        c.execute(f"ALTER TABLE quizzes ADD COLUMN {col} TEXT DEFAULT ''")
                    print(f"  ✅ Đã thêm cột {col}")
                except Exception as e:
                    print(f"  ❌ Lỗi thêm cột {col}: {e}")
        
        conn.commit()
        conn.close()
        print("✅ Đã kiểm tra và sửa database")
        
    except Exception as e:
        print(f"❌ Lỗi kiểm tra database: {e}")

def check_and_fix_questions_table():
    """Tự động kiểm tra và sửa lỗi bảng questions"""
    try:
        conn = sqlite3.connect('quiz_system.db')
        c = conn.cursor()
        
        # Kiểm tra bảng questions có tồn tại không
        c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='questions'")
        if not c.fetchone():
            print("⚠️ Bảng questions không tồn tại, đang tạo...")
            reset_questions_table()
            conn.close()
            return
        
        # Kiểm tra các cột
        c.execute("PRAGMA table_info(questions)")
        columns = [col[1] for col in c.fetchall()]
        
        required_columns = [
            'id', 'quiz_id', 'question_text', 'option_a', 'option_b', 
            'option_c', 'option_d', 'correct_answer', 'explanation',
            'question_type', 'difficulty'
        ]
        
        missing_columns = [col for col in required_columns if col not in columns]
        
        if missing_columns:
            print(f"⚠️ Thiếu cột trong questions: {missing_columns}")
            
            # Thêm các cột bị thiếu
            for col in missing_columns:
                try:
                    if col == 'difficulty':
                        c.execute(f"ALTER TABLE questions ADD COLUMN {col} TEXT DEFAULT 'medium'")
                    elif col == 'question_type':
                        c.execute(f"ALTER TABLE questions ADD COLUMN {col} TEXT DEFAULT 'multiple_choice'")
                    elif col == 'quiz_id':
                        c.execute(f"ALTER TABLE questions ADD COLUMN {col} INTEGER")
                    elif col in ['option_a', 'option_b', 'option_c', 'option_d', 'question_text', 'correct_answer', 'explanation']:
                        c.execute(f"ALTER TABLE questions ADD COLUMN {col} TEXT DEFAULT ''")
                    elif col == 'id':
                        # Đây là primary key, nếu thiếu thì phải tạo lại bảng
                        print(f"  ⚠️ Cột {col} là primary key, cần tạo lại bảng")
                        reset_questions_table()
                        break
                    else:
                        c.execute(f"ALTER TABLE questions ADD COLUMN {col} TEXT DEFAULT ''")
                    print(f"  ✅ Đã thêm cột {col}")
                except Exception as e:
                    print(f"  ❌ Lỗi thêm cột {col}: {e}")
        
        conn.commit()
        conn.close()
        print("✅ Đã kiểm tra và sửa bảng questions")
        
    except Exception as e:
        print(f"❌ Lỗi kiểm tra bảng questions: {e}")

def reset_questions_table():
    """Đặt lại bảng questions nếu có lỗi cấu trúc"""
    try:
        conn = sqlite3.connect('quiz_system.db')
        c = conn.cursor()
        
        # Xóa bảng cũ nếu tồn tại
        c.execute('DROP TABLE IF EXISTS questions')
        
        # Tạo bảng mới với đầy đủ cột
        c.execute('''CREATE TABLE questions
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      quiz_id INTEGER,
                      question_text TEXT,
                      option_a TEXT,
                      option_b TEXT,
                      option_c TEXT,
                      option_d TEXT,
                      correct_answer TEXT,
                      explanation TEXT,
                      question_type TEXT DEFAULT 'multiple_choice',
                      difficulty TEXT DEFAULT 'medium',
                      FOREIGN KEY (quiz_id) REFERENCES quizzes(id))''')
        
        conn.commit()
        conn.close()
        print("✅ Đã reset bảng questions thành công!")
        return True
    except Exception as e:
        print(f"❌ Lỗi reset bảng questions: {e}")
        return False

def reset_quizzes_table():
    """Đặt lại bảng quizzes nếu cần"""
    try:
        conn = sqlite3.connect('quiz_system.db')
        c = conn.cursor()
        c.execute('DROP TABLE IF EXISTS quizzes')
        c.execute('''CREATE TABLE quizzes
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      quiz_code TEXT UNIQUE,
                      title TEXT,
                      subject TEXT DEFAULT 'Lịch Sử',
                      created_at TIMESTAMP,
                      question_count INTEGER,
                      is_active BOOLEAN DEFAULT 1,
                      difficulty TEXT DEFAULT 'medium')''')
        conn.commit()
        conn.close()
        print("✅ Đã reset bảng quizzes thành công!")
        return True
    except Exception as e:
        print(f"❌ Lỗi reset bảng quizzes: {e}")
        return False

# ==================== KHỞI TẠO DATABASE ====================
def init_db():
    conn = sqlite3.connect('quiz_system.db')
    c = conn.cursor()
    
    # Bảng quizzes - CẬP NHẬT
    c.execute('''CREATE TABLE IF NOT EXISTS quizzes
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  quiz_code TEXT UNIQUE,
                  title TEXT,
                  subject TEXT DEFAULT 'Lịch Sử',
                  created_at TIMESTAMP,
                  question_count INTEGER,
                  is_active BOOLEAN DEFAULT 1,
                  difficulty TEXT DEFAULT 'medium')''')
    
    # Bảng questions
    c.execute('''CREATE TABLE IF NOT EXISTS questions
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  quiz_id INTEGER,
                  question_text TEXT,
                  option_a TEXT,
                  option_b TEXT,
                  option_c TEXT,
                  option_d TEXT,
                  correct_answer TEXT,
                  explanation TEXT,
                  question_type TEXT DEFAULT 'multiple_choice',
                  difficulty TEXT DEFAULT 'medium',
                  FOREIGN KEY (quiz_id) REFERENCES quizzes(id))''')
    
    # Bảng students (lưu thông tin học sinh)
    c.execute('''CREATE TABLE IF NOT EXISTS students
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  student_name TEXT NOT NULL,
                  class_name TEXT NOT NULL,
                  student_id TEXT UNIQUE,
                  email TEXT,
                  phone TEXT,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    conn.commit()
    conn.close()
    print("✅ Database đã được khởi tạo/kiểm tra")

# Chạy migration và init
migrate_database()
init_db()
check_and_fix_database()
check_and_fix_questions_table() 

# ==================== KHỞI TẠO GEMINI AI ====================
@st.cache_resource
def init_ai_model():
    try:
        # Lấy API key từ nhiều nguồn
        api_key = None
        
        # 1. Từ Streamlit secrets
        try:
            if hasattr(st, 'secrets'):
                api_key = st.secrets.get("GEMINI_API_KEY")
        except:
            pass
        
        # 2. Từ biến môi trường
        if not api_key:
            api_key = os.environ.get("GEMINI_API_KEY")
        
        # 3. Từ key trực tiếp
        if not api_key:
            api_key = "AIzaSyAXneM58drczCgMfm-Ihx0mzxIpiy8TmvQ"
        
        if not api_key or api_key == "your_api_key_here":
            st.warning("⚠️ Chưa cấu hình Gemini API Key")
            return None
        
        # Configure với API key
        genai.configure(api_key=api_key)
        
        # DÙNG MODEL GEMMA 3-4B
        model_name = 'gemini-1.5-flash'
        
        print(f"🤖 Đang khởi tạo model: {model_name}")
        
        # Tạo model
        model = genai.GenerativeModel(model_name)
        
        # Test ngắn
        test_response = model.generate_content(
            "Xin chào",
            generation_config={"max_output_tokens": 5}
        )
        
        if test_response.text:
            print(f"✅ AI Model đã sẵn sàng: {model_name}")
            return model
        else:
            print("❌ Model không trả về kết quả")
            return None
            
    except Exception as e:
        print(f"❌ Lỗi khởi tạo AI Model: {str(e)[:200]}")
        return None

# Khởi tạo Gemini model
gemini_model = init_ai_model()

# ==================== HÀM HELPER ====================
def extract_text_from_file(uploaded_file):
    """Trích xuất text từ file upload - PHIÊN BẢN CẢI TIẾN"""
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    try:
        # Reset file pointer về đầu
        uploaded_file.seek(0)
        
        if file_type == 'txt':
            content = uploaded_file.read()
            # Thử decode với UTF-8, nếu lỗi thì thử với ISO-8859-1
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return content.decode('utf-8-sig')
                except:
                    return content.decode('latin-1', errors='ignore')
        
        elif file_type == 'pdf':
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        
        elif file_type == 'docx':
            # PHẦN QUAN TRỌNG: Đọc DOCX đúng cách
            import docx
            
            # Lưu file tạm thời hoặc đọc từ bytes
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            text = ""
            
            # Đọc tất cả các paragraph
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Đọc cả text trong tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # Reset file pointer
            uploaded_file.seek(0)
            
            # DEBUG: In độ dài text để kiểm tra
            print(f"DEBUG: Đã đọc {len(text)} ký tự từ file DOCX")
            
            return text
        
    except Exception as e:
        print(f"❌ Lỗi đọc file {uploaded_file.name}: {e}")
        
        # Thử phương pháp dự phòng cho DOCX
        if file_type == 'docx':
            try:
                # Thử đọc như file zip (DOCX thực chất là zip)
                import zipfile
                import xml.etree.ElementTree as ET
                
                uploaded_file.seek(0)
                zip_data = io.BytesIO(uploaded_file.read())
                
                with zipfile.ZipFile(zip_data) as docx_zip:
                    # Đọc file document.xml
                    xml_content = docx_zip.read('word/document.xml')
                    
                    # Parse XML đơn giản
                    root = ET.fromstring(xml_content)
                    
                    # Lấy tất cả text
                    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    text_elements = root.findall('.//w:t', namespaces)
                    
                    text = ' '.join([elem.text for elem in text_elements if elem.text])
                    
                    print(f"DEBUG (dự phòng): Đã đọc {len(text)} ký tự từ DOCX XML")
                    return text
                    
            except Exception as e2:
                print(f"❌ Lỗi dự phòng DOCX: {e2}")
        
        return f"[File: {uploaded_file.name}] - Lỗi đọc nội dung: {str(e)[:100]}"

def debug_file_content(uploaded_file):
    """Debug nội dung file để tìm lỗi"""
    print(f"=== DEBUG FILE {uploaded_file.name} ===")
    print(f"Kích thước: {uploaded_file.size} bytes")
    print(f"Loại file: {uploaded_file.type}")
    
    # Đọc raw bytes
    uploaded_file.seek(0)
    raw_bytes = uploaded_file.read()
    print(f"Số bytes: {len(raw_bytes)}")
    
    # Hiển thị 200 bytes đầu tiên
    print(f"200 bytes đầu: {raw_bytes[:200]}")
    
    uploaded_file.seek(0)  # Reset
    return len(raw_bytes)

def get_sample_questions():
    """Câu hỏi mẫu khi không thể tạo bằng AI"""
    return {
        "questions": [
            {
                "question": "Chiến thắng Điện Biên Phủ diễn ra vào năm nào?",
                "options": {
                    "A": "1953",
                    "B": "1954",
                    "C": "1975",
                    "D": "1945"
                },
                "correct_answer": "B",
                "explanation": "Chiến dịch Điện Biên Phủ kết thúc thắng lợi vào ngày 7/5/1954, đánh dấu thắng lợi quyết định của quân dân Việt Nam trong kháng chiến chống Pháp."
            },
            {
                "question": "Ai là tác giả của Bản Tuyên ngôn Độc lập 2/9/1945?",
                "options": {
                    "A": "Hồ Chí Minh",
                    "B": "Trường Chinh",
                    "C": "Phạm Văn Đồng",
                    "D": "Võ Nguyên Giáp"
                },
                "correct_answer": "A",
                "explanation": "Chủ tịch Hồ Chí Minh đọc bản Tuyên ngôn Độc lập tại Quảng trường Ba Đình, Hà Nội, khai sinh nước Việt Nam Dân chủ Cộng hòa."
            },
            {
                "question": "Vua nào dựng nước Văn Lang - nhà nước đầu tiên của Việt Nam?",
                "options": {
                    "A": "An Dương Vương",
                    "B": "Vua Hùng",
                    "C": "Lý Thái Tổ",
                    "D": "Quang Trung"
                },
                "correct_answer": "B",
                "explanation": "Các Vua Hùng là những người có công dựng nước Văn Lang, đặt nền móng cho sự hình thành và phát triển của dân tộc Việt Nam."
            }
        ]
    }

def generate_quiz_questions_gemini(text, num_questions=5):
    """Tạo câu hỏi bằng Gemini API"""
    if not gemini_model:
        return None
    
    try:
        text = text[:3000]
        
        prompt = f"""Bạn là giáo viên lịch sử xuất sắc. Tạo {num_questions} câu hỏi trắc nghiệm từ tài liệu sau:

{text}

YÊU CẦU:
1. Tạo {num_questions} câu hỏi TRẮC NGHIỆM 4 lựa chọn (A, B, C, D)
2. Chỉ MỘT đáp án đúng duy nhất
3. Mỗi câu hỏi phải có giải thích ngắn gọn
4. Câu hỏi phải đa dạng: sự kiện, nhân vật, niên đại, địa điểm

ĐỊNH DẠNG JSON:
{{
  "questions": [
    {{
      "question": "Câu hỏi 1",
      "options": {{
        "A": "Đáp án A",
        "B": "Đáp án B",
        "C": "Đáp án C", 
        "D": "Đáp án D"
      }},
      "correct_answer": "A",
      "explanation": "Giải thích tại sao A đúng"
    }}
  ]
}}

Chỉ trả về JSON, không thêm bất kỳ text nào khác."""
        
        response = gemini_model.generate_content(
            prompt,
            generation_config={
                "max_output_tokens": 2000,
                "temperature": 0.7,
                "top_p": 0.8
            }
        )
        
        if not response.text:
            return None
            
        result_text = response.text.strip()
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        
        # Tìm JSON trong response
        json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if not json_match:
            return None
            
        quiz_data = json.loads(json_match.group())
        
        if "questions" not in quiz_data:
            return None
            
        # Validate và fix dữ liệu
        valid_questions = []
        for q in quiz_data["questions"]:
            if not isinstance(q, dict):
                continue
                
            # Đảm bảo có đủ các trường
            if "question" not in q or not q["question"].strip():
                continue
                
            if "options" not in q or not isinstance(q["options"], dict):
                continue
                
            # Đảm bảo có đủ 4 đáp án
            for key in ["A", "B", "C", "D"]:
                if key not in q["options"]:
                    q["options"][key] = f"Đáp án {key}"
            
            if "correct_answer" not in q or q["correct_answer"] not in ["A", "B", "C", "D"]:
                q["correct_answer"] = "A"
            
            if "explanation" not in q:
                q["explanation"] = "Không có giải thích"
            
            valid_questions.append(q)
        
        return {"questions": valid_questions[:num_questions]}
            
    except Exception as e:
        print(f"❌ Lỗi Gemini: {e}")
        return None

def generate_quiz_questions(text, num_questions=5):
    """Tổng hợp: Thử Gemini trước, nếu không được thì dùng câu hỏi mẫu"""
    if len(text.strip()) < 50:
        sample = get_sample_questions()
        sample["questions"] = sample["questions"][:min(num_questions, len(sample["questions"]))]
        return sample
    
    gemini_result = generate_quiz_questions_gemini(text, num_questions)
    
    if gemini_result and "questions" in gemini_result and len(gemini_result["questions"]) > 0:
        print(f"✅ AI đã tạo {len(gemini_result['questions'])} câu hỏi")
        return gemini_result
    
    sample = get_sample_questions()
    sample["questions"] = sample["questions"][:min(num_questions, len(sample["questions"]))]
    return sample

def calculate_grade(percentage):
    """Tính điểm chữ"""
    if percentage >= 90:
        return "A+", "🏆 Xuất sắc!", "#FFD700"
    elif percentage >= 80:
        return "A", "🎉 Giỏi!", "#C0C0C0"
    elif percentage >= 70:
        return "B", "👍 Khá!", "#CD7F32"
    elif percentage >= 60:
        return "C", "📚 Trung bình khá", "#4CAF50"
    elif percentage >= 50:
        return "D", "💪 Trung bình", "#FF9800"
    else:
        return "F", "🔄 Cần cố gắng hơn", "#F44336"

def validate_student_info(student_name, class_name):
    """Kiểm tra thông tin học sinh"""
    errors = []
    
    # Kiểm tra tên không rỗng
    if not student_name or len(student_name.strip()) < 3:
        errors.append("❌ **Họ tên quá ngắn!** Vui lòng nhập đầy đủ họ và tên (ít nhất 3 ký tự).")
    
    # Kiểm tra lớp không rỗng
    if not class_name or len(class_name.strip()) < 2:
        errors.append("❌ **Tên lớp không hợp lệ!** Vui lòng nhập tên lớp (ví dụ: 10A1, 11B2).")
    
    # Kiểm tra định dạng tên (có ít nhất 2 từ)
    name_parts = student_name.strip().split()
    if len(name_parts) < 2:
        errors.append("❌ **Vui lòng nhập cả họ và tên** (ví dụ: Nguyễn Văn A, Trần Thị B).")
    
    # Kiểm tra tên không chứa ký tự đặc biệt
    if re.search(r'[!@#$%^&*()_+={}\[\]:;"\'<>,.?/~`|\\]', student_name):
        errors.append("❌ **Tên không được chứa ký tự đặc biệt!**")
    
    return errors

def get_distinct_classes():
    """Lấy danh sách lớp duy nhất từ database"""
    try:
        conn = sqlite3.connect('quiz_system.db')
        c = conn.cursor()
        c.execute("SELECT DISTINCT class_name FROM results WHERE class_name != '' ORDER BY class_name")
        classes = [row[0] for row in c.fetchall()]
        conn.close()
        return classes
    except:
        return []

def get_students_by_class(class_name):
    """Lấy danh sách học sinh theo lớp"""
    try:
        conn = sqlite3.connect('quiz_system.db')
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute('''
            SELECT DISTINCT student_name, class_name, student_id 
            FROM results 
            WHERE class_name = ? 
            ORDER BY student_name
        ''', (class_name,))
        students = c.fetchall()
        conn.close()
        return students
    except:
        return []

# ==================== GIAO DIỆN CHÍNH ====================
def main():
    st.markdown('<h1 class="main-header">📚 HỆ THỐNG QUIZ LỊCH SỬ - QUẢN LÝ LỚP HỌC</h1>', unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.image("https://cdn-icons-png.flaticon.com/512/2237/2237288.png", width=100)
        st.title("🎮 MENU CHÍNH")
        
        # Các menu chính (luôn hiển thị)
        menu_options = [
            "🏠 TRANG CHỦ",
            "📤 TẠO QUIZ MỚI",
            "🎯 THAM GIA QUIZ",
            "🔍 TRA CỨU KẾT QUẢ",
            "📊 THỐNG KÊ CHI TIẾT",
            "🏆 BẢNG XẾP HẠNG",
            "📥 XUẤT BÁO CÁO"
        ]
        
        # Nếu đã đăng nhập admin, thêm menu quản lý
        if st.session_state.logged_in:
            menu_options.append("🔧 QUẢN LÝ QUIZ")
        
        menu = st.radio("CHỌN CHỨC NĂNG:", menu_options)
        
        st.markdown("---")
        
        # Hiển thị thông tin AI
        if gemini_model:
            st.success("**🤖 GEMINI AI:** ĐÃ KẾT NỐI")
            st.caption("Sẵn sàng tạo câu hỏi thông minh")
        else:
            st.warning("**⚠️ GEMINI AI:** CHƯA KẾT NỐI")
            st.caption("Đang dùng câu hỏi mẫu")
        
        st.markdown("---")
        
        # Thông tin nhanh
        try:
            conn = sqlite3.connect('quiz_system.db')
            c = conn.cursor()
            
            c.execute("SELECT COUNT(*) FROM quizzes")
            quiz_count = c.fetchone()[0] or 0
            
            c.execute("SELECT COUNT(DISTINCT student_name) FROM results")
            student_count_result = c.fetchone()
            student_count = student_count_result[0] if student_count_result else 0
            
            c.execute("SELECT COUNT(*) FROM results")
            test_count_result = c.fetchone()
            test_count = test_count_result[0] if test_count_result else 0
            
            c.execute("SELECT COUNT(DISTINCT class_name) FROM results WHERE class_name != ''")
            class_count_result = c.fetchone()
            class_count = class_count_result[0] if class_count_result else 0
            
            # Lấy quiz mới nhất
            c.execute("SELECT quiz_code FROM quizzes ORDER BY created_at DESC LIMIT 1")
            latest_quiz = c.fetchone()
            latest_quiz_code = latest_quiz[0] if latest_quiz else "Chưa có"
            
            conn.close()
            
            if student_count == 0:
                st.info(f"""
                **📊 THỐNG KÊ NHANH:**
                - 📝 **Quiz mới nhất:** {latest_quiz_code}
                - ⚠️ **Học sinh:** {student_count} (Chưa có ai tham gia)
                - 🏫 **Lớp học:** {class_count}
                - 📋 **Bài thi:** {test_count}
                """)
                st.warning("Chưa có học sinh tham gia. Hãy chia sẻ mã Quiz!")
            else:
                st.info(f"""
                **📊 THỐNG KÊ NHANH:**
                - 📝 **Quiz:** {quiz_count}
                - 👨‍🎓 **Học sinh:** {student_count}
                - 🏫 **Lớp học:** {class_count}
                - 📋 **Bài thi:** {test_count}
                """)
                
        except Exception as e:
            st.info("📊 Đang khởi tạo hệ thống...")
        
        st.markdown("---")
        st.caption("© 2024 Hệ thống Quiz Lịch Sử")
    
    # Gọi form đăng nhập admin (hiển thị ở sidebar)
    login_admin()
    
    # ==================== TRANG CHỦ ====================
    if menu == "🏠 TRANG CHỦ":
        st.success("🎉 **CHÀO MỪNG ĐẾN VỚI HỆ THỐNG QUIZ LỊCH SỬ THÔNG MINH**")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("""
            ### ✨ **TÍNH NĂNG NỔI BẬT:**
            
            **📚 TẠO QUIZ THÔNG MINH:**
            - 🤖 AI tự động tạo câu hỏi từ giáo án
            - 📤 Hỗ trợ đa định dạng: TXT, PDF, DOCX
            - 🎯 Tùy chỉnh số câu hỏi, độ khó
            
            **🎯 THAM GIA QUIZ:**
            - 👨‍🎓 **BẮT BUỘC** nhập thông tin học sinh
            - 🏫 Quản lý theo lớp học chi tiết
            - 📝 Giao diện thân thiện, dễ sử dụng
            
            **📊 THỐNG KÊ ĐẦY ĐỦ:**
            - 🔍 Tra cứu theo tên học sinh & lớp
            - 📈 Thống kê điểm số chi tiết
            - 🏆 Bảng xếp hạng toàn trường
            
            **🔍 TRA CỨU LINH HOẠT:**
            - Tìm kiếm theo tên, lớp, mã quiz
            - Lọc theo điểm số, thời gian
            - 📥 Xuất báo cáo Excel chi tiết
            
            **📱 TÍCH HỢP ĐA NỀN TẢNG:**
            - Hoạt động trên điện thoại & máy tính
            - Tự động lưu trữ & backup dữ liệu
            - Bảo mật thông tin học sinh
            """)
            
            st.markdown("""
            ### ⚠️ **QUY ĐỊNH QUAN TRỌNG:**
            1. **Học sinh BẮT BUỘC nhập đầy đủ họ tên và lớp**
            2. **Thông tin phải chính xác để tra cứu kết quả**
            3. **Không nhập thông tin sẽ KHÔNG được nộp bài**
            4. **Mỗi học sinh chỉ được làm bài 1 lần/quiz**
            """)
        
        with col2:
            st.markdown("### 🚀 **BẮT ĐẦU NHANH**")
            
            # Card hướng dẫn
            st.markdown("""
            <div class="student-info-card">
                <h4>📋 HƯỚNG DẪN SỬ DỤNG:</h4>
                <ol>
                    <li><strong>Tạo quiz</strong> từ file giáo án</li>
                    <li><strong>Chia sẻ mã quiz</strong> cho học sinh</li>
                    <li><strong>Học sinh tham gia</strong> (nhập đủ thông tin)</li>
                    <li><strong>Theo dõi kết quả</strong> real-time</li>
                    <li><strong>Xuất báo cáo</strong> Excel</li>
                </ol>
            </div>
            """, unsafe_allow_html=True)
            
            # Nút điều hướng nhanh
            if st.button("📤 TẠO QUIZ MỚI", use_container_width=True):
                st.session_state.menu = "📤 TẠO QUIZ MỚI"
                st.rerun()
            
            if st.button("🎯 THAM GIA QUIZ", use_container_width=True):
                st.session_state.menu = "🎯 THAM GIA QUIZ"
                st.rerun()
            
            if st.button("🔍 TRA CỨU KẾT QUẢ", use_container_width=True):
                st.session_state.menu = "🔍 TRA CỨU KẾT QUẢ"
                st.rerun()
            
            if st.button("📊 THỐNG KÊ", use_container_width=True):
                st.session_state.menu = "📊 THỐNG KÊ CHI TIẾT"
                st.rerun()
        
        # Hiển thị quiz mới nhất
        st.markdown("---")
        st.subheader("📝 **QUIZ MỚI NHẤT**")
        
        try:
            conn = sqlite3.connect('quiz_system.db')
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute('SELECT * FROM quizzes ORDER BY created_at DESC LIMIT 5')
            recent_quizzes = c.fetchall()
            conn.close()
            
            if recent_quizzes:
                cols = st.columns(len(recent_quizzes))
                for idx, quiz in enumerate(recent_quizzes):
                    with cols[idx]:
                        st.markdown(f"""
                        <div class="quiz-card">
                            <h4>{quiz['title'][:20]}...</h4>
                            <p><strong>Mã:</strong> {quiz['quiz_code']}</p>
                            <p><strong>Số câu:</strong> {quiz['question_count']}</p>
                            <small>{quiz['created_at'][:10]}</small>
                        </div>
                        """, unsafe_allow_html=True)
            else:
                st.info("📭 Chưa có quiz nào được tạo")
        except:
            st.info("📭 Đang tải dữ liệu...")
    
    # ==================== TẠO QUIZ MỚI ====================
    elif menu == "📤 TẠO QUIZ MỚI":
        st.header("📤 TẠO QUIZ MỚI TỪ GIÁO ÁN")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            uploaded_file = st.file_uploader(
                "**📁 CHỌN FILE GIÁO ÁN:**",
                type=['txt', 'pdf', 'docx'],
                help="Tải lên file giáo án lịch sử (TXT, PDF hoặc DOCX)"
            )
            
            if uploaded_file:
                with st.expander("👁️ **XEM TRƯỚC NỘI DUNG**", expanded=False):
                    text = extract_text_from_file(uploaded_file)
                    if len(text) > 1000:
                        st.text_area("Nội dung", text[:1000] + "...", height=200, disabled=True)
                    else:
                        st.text_area("Nội dung", text, height=200, disabled=True)
        
        with col2:
            num_questions = st.slider(
                "**SỐ CÂU HỎI:**",
                min_value=3,
                max_value=20,
                value=10,
                help="Chọn số lượng câu hỏi muốn tạo"
            )
            
            quiz_title = st.text_input(
                "**TIÊU ĐỀ QUIZ:**",
                value="Kiểm tra Lịch Sử",
                help="Đặt tên cho quiz của bạn"
            )
            
            subject = st.selectbox(
                "**MÔN HỌC:**",
                ["Lịch Sử", "Địa Lý", "Giáo Dục Công Dân", "Toán", "Ngữ Văn", "Tiếng Anh", "Vật Lý", "Hóa Học", "Sinh Học", "Khác"]
            )
            
            difficulty = st.select_slider(
                "**ĐỘ KHÓ:**",
                options=["Dễ", "Trung bình", "Khó"],
                value="Trung bình"
            )
        
        if uploaded_file and st.button("🚀 TẠO QUIZ BẰNG AI", type="primary", use_container_width=True):
            with st.spinner("🤖 **AI ĐANG TẠO CÂU HỎI...**" if gemini_model else "📝 **ĐANG TẠO QUIZ...**"):
                # Debug file
                file_size = debug_file_content(uploaded_file)
                
                text = extract_text_from_file(uploaded_file)
                
                # Hiển thị thông tin debug
                st.info(f"**Thông tin file:** {uploaded_file.name} ({file_size} bytes)")
                st.info(f"**Đã đọc được:** {len(text)} ký tự")
                
                if len(text) < 100:
                    st.error(f"❌ **CHỈ ĐỌC ĐƯỢC {len(text)} KÝ TỰ!** Có thể file bị lỗi định dạng.")
                    
                    # Hiển thị nội dung đã đọc được
                    with st.expander("📄 Xem nội dung đã đọc được"):
                        st.text(text[:500] + "..." if len(text) > 500 else text)
                else:
                    quiz_data = generate_quiz_questions(text, num_questions)
                    
                    if not quiz_data or "questions" not in quiz_data:
                        st.error("❌ **KHÔNG THỂ TẠO CÂU HỎI!** Vui lòng thử lại với file khác.")
                        st.stop()
                    
                    # Tạo mã quiz ngẫu nhiên
                    quiz_code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
                    
                    # Lưu vào database
                    conn = None
                    try:
                        conn = sqlite3.connect('quiz_system.db')
                        c = conn.cursor()
                        
                        # Tạo bảng nếu chưa tồn tại (đảm bảo chắc chắn)
                        c.execute('''CREATE TABLE IF NOT EXISTS quizzes
                                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                                      quiz_code TEXT UNIQUE,
                                      title TEXT,
                                      subject TEXT DEFAULT 'Lịch Sử',
                                      created_at TIMESTAMP,
                                      question_count INTEGER,
                                      is_active BOOLEAN DEFAULT 1,
                                      difficulty TEXT DEFAULT 'medium')''')
                        
                        c.execute('''CREATE TABLE IF NOT EXISTS questions
                                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                                      quiz_id INTEGER,
                                      question_text TEXT,
                                      option_a TEXT,
                                      option_b TEXT,
                                      option_c TEXT,
                                      option_d TEXT,
                                      correct_answer TEXT,
                                      explanation TEXT,
                                      question_type TEXT DEFAULT 'multiple_choice',
                                      difficulty TEXT DEFAULT 'medium',
                                      FOREIGN KEY (quiz_id) REFERENCES quizzes(id))''')
                        
                        conn.commit()
                        
                        # Kiểm tra xem bảng có tồn tại không
                        c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='quizzes'")
                        if not c.fetchone():
                            st.error("❌ **LỖI DATABASE!** Không thể tạo bảng quizzes.")
                            conn.close()
                            st.stop()
                        
                        # Lưu thông tin quiz
                        try:
                            c.execute('''INSERT INTO quizzes 
                                         (quiz_code, title, subject, created_at, question_count, difficulty) 
                                         VALUES (?, ?, ?, ?, ?, ?)''',
                                     (quiz_code, f"{subject} - {quiz_title}", subject, 
                                      datetime.now().isoformat(), len(quiz_data['questions']), difficulty))
                            quiz_id = c.lastrowid
                        except Exception as e:
                            st.error(f"❌ **LỖI LƯU QUIZ:** {str(e)[:100]}")
                            # Thử với cấu trúc đơn giản hơn
                            try:
                                c.execute('''INSERT INTO quizzes 
                                             (quiz_code, title, subject, created_at, question_count) 
                                             VALUES (?, ?, ?, ?, ?)''',
                                         (quiz_code, f"{subject} - {quiz_title}", subject, 
                                          datetime.now().isoformat(), len(quiz_data['questions'])))
                                quiz_id = c.lastrowid
                            except Exception as e2:
                                st.error(f"❌ **LỖI NẶNG LƯU QUIZ:** {str(e2)[:100]}")
                                conn.close()
                                st.stop()
                        
                        # Commit sau khi insert quiz
                        conn.commit()
                        
                        # Lưu các câu hỏi
                        st.write("💾 **Đang lưu câu hỏi vào database...**")
                        
                        success_count = 0
                        error_count = 0
                        
                        # Tạo progress bar cho từng câu hỏi
                        question_progress = st.progress(0)
                        
                        for idx, q in enumerate(quiz_data['questions']):
                            try:
                                # Hiển thị tiến trình
                                progress_percent = (idx + 1) / len(quiz_data['questions'])
                                question_progress.progress(progress_percent)
                                
                                # Đảm bảo tất cả trường đều có giá trị
                                question_text = str(q.get('question', f"Câu hỏi {idx+1}"))[:500]
                                
                                # Xử lý options
                                options = q.get('options', {})
                                option_a = str(options.get('A', 'Đáp án A'))[:200]
                                option_b = str(options.get('B', 'Đáp án B'))[:200]
                                option_c = str(options.get('C', 'Đáp án C'))[:200]
                                option_d = str(options.get('D', 'Đáp án D'))[:200]
                                
                                # Xử lý đáp án đúng
                                correct_answer = str(q.get('correct_answer', 'A')).upper()[:1]
                                if correct_answer not in ['A', 'B', 'C', 'D']:
                                    correct_answer = 'A'
                                
                                # Xử lý giải thích
                                explanation = str(q.get('explanation', 'Không có giải thích'))[:500]
                                
                                # Thực hiện INSERT
                                c.execute('''INSERT INTO questions 
                                             (quiz_id, question_text, option_a, option_b, option_c, option_d, 
                                              correct_answer, explanation, difficulty)
                                             VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                                         (quiz_id, 
                                          question_text,
                                          option_a,
                                          option_b,
                                          option_c,
                                          option_d,
                                          correct_answer,
                                          explanation,
                                          difficulty))
                                success_count += 1
                                
                            except Exception as e:
                                error_count += 1
                                st.error(f"❌ Lỗi câu hỏi {idx+1}: {str(e)[:100]}")
                        
                        question_progress.progress(1.0)
                        
                        # Commit tất cả câu hỏi
                        conn.commit()
                        
                        # Cập nhật số câu hỏi thực tế đã lưu
                        if success_count > 0:
                            try:
                                c.execute('UPDATE quizzes SET question_count = ? WHERE id = ?', (success_count, quiz_id))
                                conn.commit()
                                st.success(f"✅ **ĐÃ LƯU THÀNH CÔNG {success_count} CÂU HỎI**")
                            except Exception as e:
                                st.warning(f"⚠️ Đã lưu câu hỏi nhưng không cập nhật được số lượng: {str(e)[:100]}")
                        else:
                            # Nếu không lưu được câu hỏi nào, xóa quiz đã tạo
                            c.execute('DELETE FROM quizzes WHERE id = ?', (quiz_id,))
                            conn.commit()
                            st.error("❌ **KHÔNG THỂ LƯU CÂU HỎI!** Quiz đã bị hủy.")
                            conn.close()
                            st.stop()
                        
                        # Đóng kết nối database
                        conn.close()
                        
                        # HIỂN THỊ KẾT QUẢ
                        if success_count > 0:
                            st.success(f"🎉 **QUIZ ĐÃ ĐƯỢC TẠO THÀNH CÔNG! ({success_count}/{len(quiz_data['questions'])} câu)**")
                            
                            col_code, col_info = st.columns(2)
                            with col_code:
                                st.markdown(f"""
                                <div class="student-info-card">
                                    <h3>📋 THÔNG TIN QUIZ</h3>
                                    <p><strong>🏷️ Tiêu đề:</strong> {quiz_title}</p>
                                    <p><strong>📚 Môn học:</strong> {subject}</p>
                                    <p><strong>📊 Độ khó:</strong> {difficulty}</p>
                                    <p><strong>🔢 Số câu:</strong> {success_count}</p>
                                </div>
                                """, unsafe_allow_html=True)
                            
                            with col_info:
                                st.markdown(f"""
                                <div class="student-info-card">
                                    <h3>🎯 MÃ QUIZ</h3>
                                    <h1 style="text-align: center; color: #3B82F6;">{quiz_code}</h1>
                                    <p style="text-align: center; font-size: 0.9em;">Chia sẻ mã này cho học sinh</p>
                                </div>
                                """, unsafe_allow_html=True)
                            
                            # Hiển thị mã quiz lớn để copy
                            st.code(quiz_code, language="text")
                            
                            # Nút copy
                            if st.button("📋 Sao chép mã quiz", key="copy_quiz_code"):
                                st.info(f"✅ Đã sao chép mã: {quiz_code}")
                            
                            # Hiển thị cảnh báo nếu có lỗi
                            if error_count > 0:
                                st.warning(f"⚠️ Có {error_count} câu hỏi không lưu được. Vui lòng kiểm tra lại!")
                            
                            # Xem trước câu hỏi
                            with st.expander("📝 **XEM TRƯỚC CÂU HỎI**", expanded=False):
                                for i, q in enumerate(quiz_data['questions'][:5]):  # Chỉ hiển thị 5 câu đầu
                                    st.markdown(f"### ❓ **Câu {i+1}:** {q.get('question', 'Không có câu hỏi')}")
                                    
                                    cols = st.columns(2)
                                    with cols[0]:
                                        st.markdown(f"**A.** {q.get('options', {}).get('A', 'Đáp án A')}")
                                        st.markdown(f"**B.** {q.get('options', {}).get('B', 'Đáp án B')}")
                                    with cols[1]:
                                        st.markdown(f"**C.** {q.get('options', {}).get('C', 'Đáp án C')}")
                                        st.markdown(f"**D.** {q.get('options', {}).get('D', 'Đáp án D')}")
                                    
                                    st.markdown(f"✅ **Đáp án đúng:** {q.get('correct_answer', 'A')}")
                                    st.markdown(f"💡 **Giải thích:** {q.get('explanation', 'Không có giải thích')}")
                                    st.markdown("---")
                                
                                if len(quiz_data['questions']) > 5:
                                    st.info(f"... và {len(quiz_data['questions']) - 5} câu hỏi khác")
                        
                    except sqlite3.Error as e:
                        if conn:
                            conn.rollback()
                            conn.close()
                        st.error(f"❌ **LỖI DATABASE NGHIÊM TRỌNG:** {str(e)[:200]}")
                        st.info("💡 **Gợi ý khắc phục:**")
                        st.markdown("""
                        1. **Reset database:** Vào trang chủ để kiểm tra lại
                        2. **Kiểm tra quyền ghi:** Đảm bảo app có quyền ghi database
                        3. **Dùng database mới:** Thử xóa file quiz_system.db để tạo lại
                        """)
                    except Exception as e:
                        if conn:
                            conn.close()
                        st.error(f"❌ **LỖI KHÔNG XÁC ĐỊNH:** {str(e)[:200]}")
    
    # ==================== THAM GIA QUIZ ====================
    elif menu == "🎯 THAM GIA QUIZ":
        st.header("🎯 THAM GIA LÀM BÀI QUIZ")
        
        tab1, tab2 = st.tabs(["📝 Làm bài mới", "📋 Xem lại bài đã làm"])
        
        with tab1:
            st.markdown("### 📋 NHẬP MÃ QUIZ")
            
            quiz_code = st.text_input(
                "**Nhập mã Quiz nhận từ giáo viên:**",
                placeholder="VD: ABC123XYZ",
                help="Nhập mã 8 ký tự mà giáo viên đã cung cấp",
                key="take_quiz_code"
            ).strip().upper()
            
            if quiz_code:
                conn = sqlite3.connect('quiz_system.db')
                conn.row_factory = sqlite3.Row
                c = conn.cursor()
                
                c.execute('SELECT * FROM quizzes WHERE quiz_code = ? AND is_active = 1', (quiz_code,))
                quiz = c.fetchone()
                
                if not quiz:
                    st.error("❌ **MÃ QUIZ KHÔNG TỒN TẠI HOẶC ĐÃ BỊ KHÓA!**")
                    conn.close()
                    st.stop()
                else:
                    st.success(f"✅ **ĐÃ TÌM THẤY QUIZ:** {quiz['title']}")
                    
                    # Lấy câu hỏi
                    c.execute('SELECT * FROM questions WHERE quiz_id = ? ORDER BY id', (quiz['id'],))
                    questions = c.fetchall()
                    conn.close()
                    
                    if not questions:
                        st.error("❌ **QUIZ NÀY CHƯA CÓ CÂU HỎI!**")
                        st.stop()
                    
                    # THÔNG TIN HỌC SINH - BẮT BUỘC NHẬP
                    st.markdown("### 👨‍🎓 **THÔNG TIN HỌC SINH (BẮT BUỘC)**")
                    
                    st.markdown("""
                    <div class="warning-card">
                        <h4>⚠️ THÔNG TIN BẮT BUỘC</h4>
                        <p>Vui lòng nhập đầy đủ thông tin để:</p>
                        <ul>
                            <li>Xác định kết quả bài thi</li>
                            <li>Tra cứu kết quả sau này</li>
                            <li>Xếp hạng chính xác</li>
                            <li>Nhận chứng chỉ/giấy khen (nếu có)</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        student_name = st.text_input(
                            "**Họ và tên:**",
                            placeholder="Nguyễn Văn A",
                            help="Nhập họ tên đầy đủ",
                            key="take_student_name"
                        )
                    
                    with col2:
                        class_name = st.text_input(
                            "**Lớp:**",
                            placeholder="10A1",
                            help="Nhập tên lớp",
                            key="take_class_name"
                        )
                    
                    with col3:
                        student_id = st.text_input(
                            "**Mã học sinh (nếu có):**",
                            placeholder="HS001",
                            help="Mã số học sinh (không bắt buộc)",
                            key="take_student_id"
                        )
                    
                    # Kiểm tra thông tin bắt buộc
                    if not student_name or not class_name:
                        st.error("""
                        ⚠️ **VUI LÒNG NHẬP ĐẦY ĐỦ THÔNG TIN TRƯỚC KHI LÀM BÀI!**
                        
                        **THÔNG TIN BẮT BUỘC:**
                        1. **👨‍🎓 Họ và tên** 
                        2. **🏫 Lớp học**
                        
                        **LƯU Ý:** Không nhập thông tin sẽ KHÔNG được làm bài!
                        """)
                        st.stop()
                    
                    # Validate thông tin
                    validation_errors = validate_student_info(student_name, class_name)
                    if validation_errors:
                        for error in validation_errors:
                            st.error(error)
                        st.stop()
                    
                    # Hiển thị thông tin đã nhập
                    st.markdown(f"""
                    <div class="student-info-card">
                        <h4>📋 THÔNG TIN BÀI THI</h4>
                        <p><strong>👨‍🎓 Học sinh:</strong> {student_name}</p>
                        <p><strong>🏫 Lớp:</strong> {class_name}</p>
                        <p><strong>🆔 Mã HS:</strong> {student_id if student_id else 'Không có'}</p>
                        <p><strong>📝 Mã Quiz:</strong> {quiz_code}</p>
                        <p><strong>🔢 Số câu:</strong> {len(questions)}</p>
                        <p><strong>⏱️ Thời gian bắt đầu:</strong> {datetime.now().strftime('%H:%M %d/%m/%Y')}</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    st.subheader("📝 **BẮT ĐẦU LÀM BÀI**")
                    
                    # Lưu câu trả lời trong session state
                    if 'quiz_answers' not in st.session_state:
                        st.session_state.quiz_answers = {}
                    
                    answers = st.session_state.quiz_answers
                    
                    for i, q in enumerate(questions):
                        st.markdown(f"### **Câu {i+1}:** {q['question_text']}")
                        
                        # Hiển thị các lựa chọn
                        options = [
                            ("A", q['option_a']),
                            ("B", q['option_b']),
                            ("C", q['option_c']),
                            ("D", q['option_d'])
                        ]
                        
                        selected = answers.get(str(q['id']))
                        
                        # Tạo các nút lựa chọn
                        cols = st.columns(4)
                        for idx, (opt_key, opt_text) in enumerate(options):
                            with cols[idx]:
                                button_text = f"{opt_key}: {opt_text[:30]}..." if len(opt_text) > 30 else f"{opt_key}: {opt_text}"
                                if st.button(
                                    button_text,
                                    key=f"opt_{q['id']}_{opt_key}",
                                    type="primary" if selected == opt_key else "secondary",
                                    use_container_width=True
                                ):
                                    answers[str(q['id'])] = opt_key
                                    st.rerun()
                        
                        # Hiển thị đã chọn
                        if selected:
                            option_texts = {
                                'A': q['option_a'],
                                'B': q['option_b'],
                                'C': q['option_c'],
                                'D': q['option_d']
                            }
                            st.info(f"✅ **Bạn đã chọn:** **{selected}** - {option_texts[selected]}")
                        
                        st.markdown("---")
                    
                    # Nút nộp bài
                    if st.button("📤 **NỘP BÀI THI**", type="primary", use_container_width=True):
                        if len(answers) < len(questions):
                            st.warning(f"⚠️ **BẠN MỚI TRẢ LỜI {len(answers)}/{len(questions)} CÂU!** Vẫn nộp bài?")
                        
                        # Tính điểm
                        score = 0
                        details = []
                        
                        for q in questions:
                            question_id = str(q['id'])
                            user_answer = answers.get(question_id, '').upper()
                            is_correct = (user_answer == q['correct_answer'])
                            
                            if is_correct:
                                score += 1
                            
                            details.append({
                                'question': q['question_text'],
                                'user_answer': user_answer if user_answer else 'Không trả lời',
                                'correct_answer': q['correct_answer'],
                                'is_correct': is_correct,
                                'explanation': q['explanation']
                            })
                        
                        # Tính phần trăm và xếp loại
                        percentage = (score / len(questions)) * 100
                        grade, evaluation, grade_color = calculate_grade(percentage)
                        
                        # Lưu kết quả
                        conn = sqlite3.connect('quiz_system.db')
                        c = conn.cursor()
                        c.execute('''INSERT INTO results 
                                     (quiz_code, student_name, class_name, student_id, 
                                      score, total_questions, percentage, grade, submitted_at)
                                     VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                                 (quiz_code, student_name, class_name, student_id,
                                  score, len(questions), percentage, grade, datetime.now()))
                        conn.commit()
                        
                        # Lấy ID kết quả vừa lưu
                        result_id = c.lastrowid
                        conn.close()
                        
                        # Hiển thị kết quả
                        st.markdown(f"""
                        <div class="score-card">
                            <h1>{evaluation.split()[-1]}</h1>
                            <h2>{evaluation}</h2>
                            <h3>Điểm: {score}/{len(questions)}</h3>
                            <p>Tỉ lệ: {percentage:.1f}% | Xếp loại: <span style="color: {grade_color}">{grade}</span></p>
                            <p><small>Mã bài thi: {result_id}</small></p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Thông tin lưu trữ
                        st.markdown(f"""
                        <div class="student-info-card">
                            <h4>✅ ĐÃ LƯU KẾT QUẢ</h4>
                            <p><strong>🆔 Mã bài thi:</strong> {result_id} (Ghi nhớ để tra cứu sau)</p>
                            <p><strong>📋 Mã Quiz:</strong> {quiz_code}</p>
                            <p><strong>👨‍🎓 Học sinh:</strong> {student_name}</p>
                            <p><strong>🏫 Lớp:</strong> {class_name}</p>
                            <p><strong>📅 Thời gian:</strong> {datetime.now().strftime('%H:%M %d/%m/%Y')}</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Chi tiết từng câu
                        with st.expander("📋 **XEM CHI TIẾT TỪNG CÂU**", expanded=False):
                            for i, detail in enumerate(details):
                                if detail['is_correct']:
                                    st.success(f"**Câu {i+1}:** {detail['question']}")
                                    st.markdown(f"✅ **Bạn chọn:** **{detail['user_answer']}** (Đúng)")
                                else:
                                    st.error(f"**Câu {i+1}:** {detail['question']}")
                                    st.markdown(f"❌ **Bạn chọn:** **{detail['user_answer']}**")
                                    st.markdown(f"✅ **Đáp án đúng:** **{detail['correct_answer']}**")
                                
                                st.markdown(f"💡 **Giải thích:** {detail['explanation']}")
                                st.markdown("---")
                        
                        # Xóa session state
                        if 'quiz_answers' in st.session_state:
                            del st.session_state.quiz_answers
                        
                        st.balloons()
                        st.info("💡 **LƯU Ý:** Ghi nhớ mã bài thi ({}) để tra cứu lại kết quả sau này!".format(result_id))
        
        with tab2:
            st.markdown("### 🔍 **TRA CỨU BÀI ĐÃ LÀM**")
            
            search_option = st.radio(
                "Tìm kiếm theo:",
                ["Tên học sinh", "Mã bài thi", "Mã Quiz"],
                horizontal=True
            )
            
            if search_option == "Tên học sinh":
                col1, col2 = st.columns(2)
                with col1:
                    search_name = st.text_input("Nhập tên học sinh:", placeholder="Nguyễn Văn A")
                with col2:
                    search_class = st.text_input("Nhập lớp:", placeholder="10A1")
                
                if search_name:
                    conn = sqlite3.connect('quiz_system.db')
                    conn.row_factory = sqlite3.Row
                    c = conn.cursor()
                    
                    if search_class:
                        c.execute('''
                            SELECT * FROM results 
                            WHERE student_name LIKE ? AND class_name LIKE ?
                            ORDER BY submitted_at DESC
                            LIMIT 20
                        ''', (f'%{search_name}%', f'%{search_class}%'))
                    else:
                        c.execute('''
                            SELECT * FROM results 
                            WHERE student_name LIKE ?
                            ORDER BY submitted_at DESC
                            LIMIT 20
                        ''', (f'%{search_name}%',))
                    
                    results = c.fetchall()
                    conn.close()
                    
                    if results:
                        st.success(f"✅ Tìm thấy {len(results)} bài thi")
                        
                        for r in results:
                            with st.expander(f"📝 {r['student_name']} - {r['class_name']} - {r['quiz_code']} ({r['submitted_at'][:16]})"):
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("Điểm", f"{r['score']}/{r['total_questions']}")
                                with col2:
                                    st.metric("Tỉ lệ", f"{r['percentage']:.1f}%")
                                with col3:
                                    grade_color = {
                                        'A+': '#FFD700', 'A': '#C0C0C0', 'B': '#CD7F32',
                                        'C': '#4CAF50', 'D': '#FF9800', 'F': '#F44336'
                                    }.get(r['grade'], '#000000')
                                    st.markdown(f"**Xếp loại:** <span style='color: {grade_color}'>{r['grade']}</span>", unsafe_allow_html=True)
                                
                                st.info(f"**Mã bài thi:** {r['id']} | **Mã Quiz:** {r['quiz_code']}")
                    else:
                        st.info("📭 Không tìm thấy bài thi nào")
            
            elif search_option == "Mã bài thi":
                result_id = st.text_input("Nhập mã bài thi:", placeholder="123")
                if result_id and result_id.isdigit():
                    conn = sqlite3.connect('quiz_system.db')
                    conn.row_factory = sqlite3.Row
                    c = conn.cursor()
                    c.execute('SELECT * FROM results WHERE id = ?', (int(result_id),))
                    result = c.fetchone()
                    conn.close()
                    
                    if result:
                        st.success("✅ Tìm thấy bài thi")
                        
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Điểm", f"{result['score']}/{result['total_questions']}")
                        with col2:
                            st.metric("Tỉ lệ", f"{result['percentage']:.1f}%")
                        with col3:
                            grade_color = {
                                'A+': '#FFD700', 'A': '#C0C0C0', 'B': '#CD7F32',
                                'C': '#4CAF50', 'D': '#FF9800', 'F': '#F44336'
                            }.get(result['grade'], '#000000')
                            st.markdown(f"**Xếp loại:** <span style='color: {grade_color}'>{result['grade']}</span>", unsafe_allow_html=True)
                        
                        st.info(f"""
                        **Thông tin chi tiết:**
                        - **Học sinh:** {result['student_name']}
                        - **Lớp:** {result['class_name']}
                        - **Mã HS:** {result['student_id'] or 'Không có'}
                        - **Mã Quiz:** {result['quiz_code']}
                        - **Thời gian:** {result['submitted_at']}
                        """)
                    else:
                        st.error("❌ Không tìm thấy bài thi với mã này")
            
            else:  # Mã Quiz
                search_quiz = st.text_input("Nhập mã Quiz:", placeholder="ABC123XYZ").upper()
                if search_quiz:
                    conn = sqlite3.connect('quiz_system.db')
                    conn.row_factory = sqlite3.Row
                    c = conn.cursor()
                    c.execute('SELECT * FROM results WHERE quiz_code = ? ORDER BY percentage DESC', (search_quiz,))
                    results = c.fetchall()
                    conn.close()
                    
                    if results:
                        st.success(f"✅ Tìm thấy {len(results)} bài thi cho Quiz {search_quiz}")
                        
                        # Hiển thị bảng xếp hạng
                        data = []
                        for i, r in enumerate(results):
                            data.append({
                                "Hạng": i+1,
                                "Họ tên": r['student_name'],
                                "Lớp": r['class_name'],
                                "Điểm": f"{r['score']}/{r['total_questions']}",
                                "Tỉ lệ": f"{r['percentage']:.1f}%",
                                "Xếp loại": r['grade'],
                                "Thời gian": r['submitted_at'][:16]
                            })
                        
                        df = pd.DataFrame(data)
                        st.dataframe(df, use_container_width=True, hide_index=True)
                    else:
                        st.info("📭 Không tìm thấy bài thi nào cho Quiz này")
    
    # ==================== TRA CỨU KẾT QUẢ ====================
    elif menu == "🔍 TRA CỨU KẾT QUẢ":
        st.header("🔍 TRA CỨU KẾT QUẢ CHI TIẾT")
        
        st.markdown("""
        <div class="student-info-card">
            <h4>🎯 TÌM KIẾM THEO NHIỀU TIÊU CHÍ</h4>
            <p>Tìm kiếm linh hoạt theo tên học sinh, lớp, mã quiz, hoặc khoảng điểm</p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            search_type = st.selectbox(
                "Tiêu chí chính:",
                ["Tên học sinh", "Lớp", "Mã Quiz", "Khoảng điểm"]
            )
        
        with col2:
            if search_type == "Tên học sinh":
                search_value = st.text_input("Nhập tên học sinh:", placeholder="Nguyễn Văn A")
            elif search_type == "Lớp":
                search_value = st.text_input("Nhập tên lớp:", placeholder="10A1")
            elif search_type == "Mã Quiz":
                search_value = st.text_input("Nhập mã Quiz:", placeholder="ABC123XYZ").upper()
            else:  # Khoảng điểm
                min_score = st.number_input("Điểm tối thiểu (%):", 0, 100, 0)
                max_score = st.number_input("Điểm tối đa (%):", 0, 100, 100)
        
        with col3:
            date_from = st.date_input("Từ ngày:", value=None)
            date_to = st.date_input("Đến ngày:", value=None)
            show_all = st.checkbox("Hiển thị tất cả", value=False)
        
        if st.button("🔎 **TÌM KIẾM**", type="primary", use_container_width=True) or show_all:
            conn = sqlite3.connect('quiz_system.db')
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            
            # Xây dựng query động
            query = "SELECT * FROM results WHERE 1=1"
            params = []
            
            if not show_all:
                if search_type == "Tên học sinh" and search_value:
                    query += " AND student_name LIKE ?"
                    params.append(f'%{search_value}%')
                elif search_type == "Lớp" and search_value:
                    query += " AND class_name LIKE ?"
                    params.append(f'%{search_value}%')
                elif search_type == "Mã Quiz" and search_value:
                    query += " AND quiz_code = ?"
                    params.append(search_value)
                elif search_type == "Khoảng điểm":
                    query += " AND percentage BETWEEN ? AND ?"
                    params.extend([min_score, max_score])
                
                if date_from:
                    query += " AND DATE(submitted_at) >= ?"
                    params.append(date_from.isoformat())
                
                if date_to:
                    query += " AND DATE(submitted_at) <= ?"
                    params.append(date_to.isoformat())
            
            query += " ORDER BY submitted_at DESC LIMIT 100"
            
            c.execute(query, params)
            results = c.fetchall()
            conn.close()
            
            if results:
                st.success(f"✅ **TÌM THẤY {len(results)} KẾT QUẢ**")
                
                # Tạo DataFrame để hiển thị
                data = []
                for r in results:
                    data.append({
                        "Mã bài": r['id'],
                        "Họ tên": r['student_name'],
                        "Lớp": r['class_name'],
                        "Mã HS": r['student_id'] or "",
                        "Mã Quiz": r['quiz_code'],
                        "Điểm": f"{r['score']}/{r['total_questions']}",
                        "Tỉ lệ": f"{r['percentage']:.1f}%",
                        "Xếp loại": r['grade'],
                        "Thời gian": r['submitted_at'][:16]
                    })
                
                df = pd.DataFrame(data)
                
                # Hiển thị bảng với định dạng đẹp
                st.dataframe(
                    df,
                    use_container_width=True,
                    column_config={
                        "Tỉ lệ": st.column_config.ProgressColumn(
                            "Tỉ lệ %",
                            help="Tỉ lệ điểm đạt được",
                            format="%.1f%%",
                            min_value=0,
                            max_value=100,
                        )
                    },
                    hide_index=True
                )
                
                # Thống kê nhanh
                if len(results) > 0:
                    avg_percentage = sum(r['percentage'] for r in results) / len(results)
                    max_percentage = max(r['percentage'] for r in results)
                    min_percentage = min(r['percentage'] for r in results)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("📊 Điểm TB", f"{avg_percentage:.1f}%")
                    with col2:
                        st.metric("🏆 Điểm cao nhất", f"{max_percentage:.1f}%")
                    with col3:
                        st.metric("📉 Điểm thấp nhất", f"{min_percentage:.1f}%")
                    with col4:
                        st.metric("📋 Số bài", len(results))
                
                # Nút xuất Excel
                excel_buffer = io.BytesIO()
                df.to_excel(excel_buffer, index=False, engine='openpyxl')
                excel_buffer.seek(0)
                
                st.download_button(
                    label="📥 **TẢI KẾT QUẢ EXCEL**",
                    data=excel_buffer,
                    file_name=f"ket_qua_tra_cuu_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            else:
                st.info("📭 **KHÔNG TÌM THẤY KẾT QUẢ NÀO PHÙ HỢP**")
    
    # ==================== THỐNG KÊ CHI TIẾT ====================
    elif menu == "📊 THỐNG KÊ CHI TIẾT":
        st.header("📊 THỐNG KÊ & PHÂN TÍCH CHI TIẾT")
        
        tab1, tab2, tab3 = st.tabs(["📈 Tổng quan", "🏫 Theo lớp", "📝 Theo Quiz"])
        
        with tab1:
            st.markdown("### 📈 **THỐNG KÊ TỔNG QUAN HỆ THỐNG**")
            
            try:
                conn = sqlite3.connect('quiz_system.db')
                conn.row_factory = sqlite3.Row
                c = conn.cursor()
                
                # Lấy dữ liệu thống kê
                c.execute("SELECT COUNT(*) as total FROM results")
                total_tests_result = c.fetchone()
                total_tests = total_tests_result['total'] if total_tests_result else 0
                
                c.execute("SELECT COUNT(DISTINCT student_name) as total FROM results")
                total_students_result = c.fetchone()
                total_students = total_students_result['total'] if total_students_result else 0
                
                c.execute("SELECT COUNT(DISTINCT class_name) as total FROM results WHERE class_name != ''")
                total_classes_result = c.fetchone()
                total_classes = total_classes_result['total'] if total_classes_result else 0
                
                c.execute("SELECT COUNT(DISTINCT quiz_code) as total FROM results")
                total_quizzes_result = c.fetchone()
                total_quizzes = total_quizzes_result['total'] if total_quizzes_result else 0
                
                c.execute("SELECT AVG(percentage) as avg FROM results")
                avg_score_result = c.fetchone()
                avg_score = avg_score_result['avg'] if avg_score_result and avg_score_result['avg'] else 0
                
                # Hiển thị metrics
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("📋 Tổng bài thi", f"{total_tests:,}")
                with col2:
                    st.metric("👨‍🎓 Tổng học sinh", f"{total_students:,}")
                with col3:
                    st.metric("🏫 Tổng lớp", f"{total_classes:,}")
                with col4:
                    st.metric("📚 Tổng Quiz", f"{total_quizzes:,}")
                
                st.metric("📊 Điểm trung bình", f"{avg_score:.1f}%", delta=f"{avg_score-50:+.1f}%")
                
                # Phân bố điểm
                st.markdown("### 📊 **PHÂN BỐ ĐIỂM SỐ**")
                c.execute('''
                    SELECT 
                        CASE 
                            WHEN percentage >= 90 THEN 'A+ (90-100%)'
                            WHEN percentage >= 80 THEN 'A (80-89%)'
                            WHEN percentage >= 70 THEN 'B (70-79%)'
                            WHEN percentage >= 60 THEN 'C (60-69%)'
                            WHEN percentage >= 50 THEN 'D (50-59%)'
                            ELSE 'F (<50%)'
                        END as grade_range,
                        COUNT(*) as count,
                        ROUND(COUNT(*) * 100.0 / (SELECT COUNT(*) FROM results), 1) as percentage
                    FROM results 
                    GROUP BY grade_range
                    ORDER BY 
                        CASE grade_range
                            WHEN 'A+ (90-100%)' THEN 1
                            WHEN 'A (80-89%)' THEN 2
                            WHEN 'B (70-79%)' THEN 3
                            WHEN 'C (60-69%)' THEN 4
                            WHEN 'D (50-59%)' THEN 5
                            ELSE 6
                        END
                ''')
                grade_dist = c.fetchall()
                
                if grade_dist:
                    # Hiển thị bảng phân bố
                    grade_data = []
                    for g in grade_dist:
                        grade_data.append({
                            "Khoảng điểm": g['grade_range'],
                            "Số bài": g['count'],
                            "Tỉ lệ": f"{g['percentage']}%"
                        })
                    
                    df_grade = pd.DataFrame(grade_data)
                    st.dataframe(df_grade, use_container_width=True, hide_index=True)
                    
                    # Hiển thị bằng metric cards
                    cols = st.columns(len(grade_dist))
                    for idx, g in enumerate(grade_dist):
                        with cols[idx]:
                            color = {
                                'A+ (90-100%)': '#FFD700',
                                'A (80-89%)': '#C0C0C0',
                                'B (70-79%)': '#CD7F32',
                                'C (60-69%)': '#4CAF50',
                                'D (50-59%)': '#FF9800',
                                'F (<50%)': '#F44336'
                            }.get(g['grade_range'], '#9E9E9E')
                            
                            st.markdown(f"""
                            <div style="text-align: center; padding: 10px; background-color: {color}; border-radius: 10px;">
                                <h4>{g['grade_range'].split()[0]}</h4>
                                <h3>{g['count']}</h3>
                                <p>{g['percentage']}%</p>
                            </div>
                            """, unsafe_allow_html=True)
                
                # Top 10 học sinh xuất sắc
                st.markdown("### 🏆 **TOP 10 HỌC SINH XUẤT SẮC**")
                c.execute('''
                    SELECT student_name, class_name,
                           COUNT(*) as test_count,
                           ROUND(AVG(percentage), 1) as avg_score,
                           MAX(percentage) as best_score
                    FROM results 
                    GROUP BY student_name, class_name
                    HAVING COUNT(*) >= 2
                    ORDER BY avg_score DESC
                    LIMIT 10
                ''')
                top_students = c.fetchall()
                
                if top_students:
                    top_data = []
                    for i, s in enumerate(top_students):
                        top_data.append({
                            "Hạng": i+1,
                            "Họ tên": s['student_name'],
                            "Lớp": s['class_name'],
                            "Số bài": s['test_count'],
                            "Điểm TB": f"{s['avg_score']}%",
                            "Điểm cao nhất": f"{s['best_score']}%"
                        })
                    
                    df_top = pd.DataFrame(top_data)
                    st.dataframe(df_top, use_container_width=True, hide_index=True)
                else:
                    st.info("📭 Chưa có đủ dữ liệu để xếp hạng")
                
                conn.close()
                
            except Exception as e:
                st.error(f"❌ Lỗi thống kê: {str(e)}")
        
        with tab2:
            st.markdown("### 🏫 **THỐNG KÊ THEO LỚP**")
            
            # Lấy danh sách lớp
            classes = get_distinct_classes()
            
            if classes:
                selected_class = st.selectbox("Chọn lớp để xem thống kê:", classes)
                
                if selected_class:
                    try:
                        conn = sqlite3.connect('quiz_system.db')
                        conn.row_factory = sqlite3.Row
                        c = conn.cursor()
                        
                        # Thống kê lớp
                        c.execute('''
                            SELECT COUNT(*) as total_tests,
                                   COUNT(DISTINCT student_name) as total_students,
                                   AVG(percentage) as avg_score,
                                   MAX(percentage) as max_score,
                                   MIN(percentage) as min_score
                            FROM results 
                            WHERE class_name = ?
                        ''', (selected_class,))
                        
                        class_stats = c.fetchone()
                        
                        # Top học sinh trong lớp
                        c.execute('''
                            SELECT student_name,
                                   COUNT(*) as test_count,
                                   ROUND(AVG(percentage), 1) as avg_score,
                                   MAX(percentage) as best_score
                            FROM results 
                            WHERE class_name = ?
                            GROUP BY student_name
                            ORDER BY avg_score DESC
                            LIMIT 10
                        ''', (selected_class,))
                        
                        top_in_class = c.fetchall()
                        
                        conn.close()
                        
                        # Hiển thị thống kê lớp
                        if class_stats['total_tests'] > 0:
                            col1, col2, col3, col4 = st.columns(4)
                            with col1:
                                st.metric("📋 Tổng bài thi", class_stats['total_tests'])
                            with col2:
                                st.metric("👨‍🎓 Số học sinh", class_stats['total_students'])
                            with col3:
                                st.metric("📊 Điểm TB", f"{class_stats['avg_score']:.1f}%")
                            with col4:
                                st.metric("📈 Điểm cao nhất", f"{class_stats['max_score']:.1f}%")
                            
                            # Danh sách học sinh trong lớp
                            st.markdown("### 👨‍🎓 **DANH SÁCH HỌC SINH TRONG LỚP**")
                            students_in_class = get_students_by_class(selected_class)
                            
                            if students_in_class:
                                student_data = []
                                for s in students_in_class:
                                    student_data.append({
                                        "Họ tên": s['student_name'],
                                        "Mã HS": s['student_id'] or "",
                                        "Lớp": s['class_name']
                                    })
                                
                                df_students = pd.DataFrame(student_data)
                                st.dataframe(df_students, use_container_width=True, hide_index=True)
                            
                            # Top học sinh trong lớp
                            if top_in_class:
                                st.markdown("### 🏆 **TOP HỌC SINH TRONG LỚP**")
                                top_data = []
                                for i, s in enumerate(top_in_class):
                                    top_data.append({
                                        "Hạng": i+1,
                                        "Họ tên": s['student_name'],
                                        "Số bài": s['test_count'],
                                        "Điểm TB": f"{s['avg_score']}%",
                                        "Điểm cao nhất": f"{s['best_score']}%"
                                    })
                                
                                df_top_class = pd.DataFrame(top_data)
                                st.dataframe(df_top_class, use_container_width=True, hide_index=True)
                        else:
                            st.info(f"📭 Lớp {selected_class} chưa có bài thi nào")
                            
                    except Exception as e:
                        st.error(f"❌ Lỗi: {str(e)}")
            else:
                st.info("📭 Chưa có dữ liệu lớp học")
        
        with tab3:
            st.markdown("### 📝 **THỐNG KÊ THEO QUIZ**")
            
            try:
                conn = sqlite3.connect('quiz_system.db')
                c = conn.cursor()
                c.execute("SELECT DISTINCT quiz_code FROM results ORDER BY quiz_code")
                quizzes = [row[0] for row in c.fetchall()]
                conn.close()
                
                if quizzes:
                    selected_quiz = st.selectbox("Chọn Quiz để xem thống kê:", quizzes)
                    
                    if selected_quiz:
                        conn = sqlite3.connect('quiz_system.db')
                        conn.row_factory = sqlite3.Row
                        c = conn.cursor()
                        
                        # Thống kê quiz
                        c.execute('''
                            SELECT COUNT(*) as total_tests,
                                   COUNT(DISTINCT student_name) as total_students,
                                   COUNT(DISTINCT class_name) as total_classes,
                                   AVG(percentage) as avg_score,
                                   MAX(percentage) as max_score,
                                   MIN(percentage) as min_score
                            FROM results 
                            WHERE quiz_code = ?
                        ''', (selected_quiz,))
                        
                        quiz_stats = c.fetchone()
                        
                        # Top học sinh trong quiz
                        c.execute('''
                            SELECT student_name, class_name,
                                   score, total_questions,
                                   percentage, grade, submitted_at
                            FROM results 
                            WHERE quiz_code = ?
                            ORDER BY percentage DESC
                            LIMIT 10
                        ''', (selected_quiz,))
                        
                        top_in_quiz = c.fetchall()
                        
                        conn.close()
                        
                        # Hiển thị thống kê quiz
                        if quiz_stats['total_tests'] > 0:
                            col1, col2, col3, col4 = st.columns(4)
                            with col1:
                                st.metric("📋 Tổng bài thi", quiz_stats['total_tests'])
                            with col2:
                                st.metric("👨‍🎓 Số học sinh", quiz_stats['total_students'])
                            with col3:
                                st.metric("🏫 Số lớp", quiz_stats['total_classes'])
                            with col4:
                                st.metric("📊 Điểm TB", f"{quiz_stats['avg_score']:.1f}%")
                            
                            # Top học sinh trong quiz
                            if top_in_quiz:
                                st.markdown("### 🏆 **TOP HỌC SINH TRONG QUIZ**")
                                top_data = []
                                for i, s in enumerate(top_in_quiz):
                                    top_data.append({
                                        "Hạng": i+1,
                                        "Họ tên": s['student_name'],
                                        "Lớp": s['class_name'],
                                        "Điểm": f"{s['score']}/{s['total_questions']}",
                                        "Tỉ lệ": f"{s['percentage']:.1f}%",
                                        "Xếp loại": s['grade']
                                    })
                                
                                df_top_quiz = pd.DataFrame(top_data)
                                st.dataframe(df_top_quiz, use_container_width=True, hide_index=True)
                        else:
                            st.info(f"📭 Quiz {selected_quiz} chưa có bài thi nào")
                            
                else:
                    st.info("📭 Chưa có dữ liệu quiz")
                    
            except Exception as e:
                st.error(f"❌ Lỗi: {str(e)}")
    
    # ==================== BẢNG XẾP HẠNG ====================
    elif menu == "🏆 BẢNG XẾP HẠNG":
        st.header("🏆 BẢNG XẾP HẠNG TOÀN TRƯỜNG")
        
        rank_type = st.radio(
            "Xếp hạng theo:",
            ["📊 Toàn trường", "🏫 Theo lớp", "📝 Theo Quiz"],
            horizontal=True
        )
        
        if rank_type == "🏫 Theo lớp":
            classes = get_distinct_classes()
            
            if classes:
                selected_class = st.selectbox("Chọn lớp:", classes)
                
                if selected_class:
                    conn = sqlite3.connect('quiz_system.db')
                    conn.row_factory = sqlite3.Row
                    c = conn.cursor()
                    c.execute('''
                        SELECT student_name, class_name, quiz_code, 
                               score, total_questions, percentage, grade, submitted_at
                        FROM results 
                        WHERE class_name = ? 
                        ORDER BY percentage DESC, submitted_at 
                        LIMIT 20
                    ''', (selected_class,))
                    rankings = c.fetchall()
                    conn.close()
                    
                    if rankings:
                        st.success(f"🏫 **BẢNG XẾP HẠNG LỚP {selected_class}**")
                        
                        for i, r in enumerate(rankings):
                            if i == 0:
                                medal = "🥇"
                                rank_class = "rank-1"
                            elif i == 1:
                                medal = "🥈"
                                rank_class = "rank-2"
                            elif i == 2:
                                medal = "🥉"
                                rank_class = "rank-3"
                            else:
                                medal = f"#{i+1}"
                                rank_class = ""
                            
                            st.markdown(f"""
                            <div class="quiz-card {rank_class}" style="margin: 10px 0;">
                                <div style="display: flex; justify-content: space-between; align-items: center;">
                                    <div>
                                        <h4 style="margin: 0;">{medal} {r['student_name']}</h4>
                                        <p style="margin: 5px 0 0 0; font-size: 0.9em;">
                                            {r['class_name']} | Quiz: {r['quiz_code']}
                                        </p>
                                    </div>
                                    <div style="text-align: right;">
                                        <h3 style="margin: 0; color: #3B82F6;">{r['percentage']:.1f}%</h3>
                                        <p style="margin: 5px 0 0 0; font-size: 0.9em;">
                                            {r['score']}/{r['total_questions']} | {r['grade']}
                                        </p>
                                    </div>
                                </div>
                                <p style="margin: 10px 0 0 0; font-size: 0.8em; color: #666;">
                                    📅 {r['submitted_at'][:16]}
                                </p>
                            </div>
                            """, unsafe_allow_html=True)
            else:
                st.info("📭 Chưa có dữ liệu lớp học")
        
        elif rank_type == "📝 Theo Quiz":
            try:
                conn = sqlite3.connect('quiz_system.db')
                c = conn.cursor()
                c.execute("SELECT DISTINCT quiz_code FROM results ORDER BY quiz_code")
                quizzes = [row[0] for row in c.fetchall()]
                conn.close()
                
                if quizzes:
                    selected_quiz = st.selectbox("Chọn mã Quiz:", quizzes)
                    
                    if selected_quiz:
                        conn = sqlite3.connect('quiz_system.db')
                        conn.row_factory = sqlite3.Row
                        c = conn.cursor()
                        c.execute('''
                            SELECT student_name, class_name, quiz_code, 
                                   score, total_questions, percentage, grade, submitted_at
                            FROM results 
                            WHERE quiz_code = ? 
                            ORDER BY percentage DESC, submitted_at 
                            LIMIT 20
                        ''', (selected_quiz,))
                        rankings = c.fetchall()
                        conn.close()
                        
                        if rankings:
                            st.success(f"📝 **BẢNG XẾP HẠNG QUIZ {selected_quiz}**")
                            
                            for i, r in enumerate(rankings):
                                if i == 0:
                                    medal = "🥇"
                                    rank_class = "rank-1"
                                elif i == 1:
                                    medal = "🥈"
                                    rank_class = "rank-2"
                                elif i == 2:
                                    medal = "🥉"
                                    rank_class = "rank-3"
                                else:
                                    medal = f"#{i+1}"
                                    rank_class = ""
                                
                                st.markdown(f"""
                                <div class="quiz-card {rank_class}" style="margin: 10px 0;">
                                    <div style="display: flex; justify-content: space-between; align-items: center;">
                                        <div>
                                            <h4 style="margin: 0;">{medal} {r['student_name']}</h4>
                                            <p style="margin: 5px 0 0 0; font-size: 0.9em;">
                                                {r['class_name']}
                                            </p>
                                        </div>
                                        <div style="text-align: right;">
                                            <h3 style="margin: 0; color: #3B82F6;">{r['percentage']:.1f}%</h3>
                                            <p style="margin: 5px 0 0 0; font-size: 0.9em;">
                                                {r['score']}/{r['total_questions']} | {r['grade']}
                                            </p>
                                        </div>
                                    </div>
                                    <p style="margin: 10px 0 0 0; font-size: 0.8em; color: #666;">
                                        📅 {r['submitted_at'][:16]}
                                    </p>
                                </div>
                                """, unsafe_allow_html=True)
                else:
                    st.info("📭 Chưa có dữ liệu quiz")
                    
            except Exception as e:
                st.error(f"❌ Lỗi: {str(e)}")
        
        else:  # Toàn trường
            try:
                conn = sqlite3.connect('quiz_system.db')
                conn.row_factory = sqlite3.Row
                c = conn.cursor()
                c.execute('''
                    SELECT student_name, class_name, quiz_code, 
                           score, total_questions, percentage, grade, submitted_at
                    FROM results 
                    ORDER BY percentage DESC, submitted_at 
                    LIMIT 20
                ''')
                rankings = c.fetchall()
                conn.close()
                
                if rankings:
                    st.success("🏆 **BẢNG XẾP HẠNG TOÀN TRƯỜNG (TOP 20)**")
                    
                    for i, r in enumerate(rankings):
                        if i == 0:
                            medal = "🥇"
                            rank_class = "rank-1"
                        elif i == 1:
                            medal = "🥈"
                            rank_class = "rank-2"
                        elif i == 2:
                            medal = "🥉"
                            rank_class = "rank-3"
                        else:
                            medal = f"#{i+1}"
                            rank_class = ""
                        
                        st.markdown(f"""
                        <div class="quiz-card {rank_class}" style="margin: 10px 0;">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <h4 style="margin: 0;">{medal} {r['student_name']}</h4>
                                    <p style="margin: 5px 0 0 0; font-size: 0.9em;">
                                        {r['class_name']} | Quiz: {r['quiz_code']}
                                    </p>
                                </div>
                                <div style="text-align: right;">
                                    <h3 style="margin: 0; color: #3B82F6;">{r['percentage']:.1f}%</h3>
                                    <p style="margin: 5px 0 0 0; font-size: 0.9em;">
                                        {r['score']}/{r['total_questions']} | {r['grade']}
                                    </p>
                                </div>
                            </div>
                            <p style="margin: 10px 0 0 0; font-size: 0.8em; color: #666;">
                                📅 {r['submitted_at'][:16]}
                            </p>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("📭 Chưa có dữ liệu xếp hạng")
                    
            except Exception as e:
                st.error(f"❌ Lỗi: {str(e)}")
    
    # ==================== QUẢN LÝ QUIZ (CHỈ ADMIN) ====================
    elif menu == "🔧 QUẢN LÝ QUIZ":
        if not st.session_state.logged_in:
            st.error("Bạn cần đăng nhập với quyền giáo viên để sử dụng tính năng này.")
            st.stop()
        
        st.header("🔧 QUẢN LÝ QUIZ (CHỈ GIÁO VIÊN)")
        
        # Lấy danh sách quiz
        conn = sqlite3.connect('quiz_system.db')
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT id, quiz_code, title, subject, question_count, created_at FROM quizzes ORDER BY created_at DESC")
        quizzes = c.fetchall()
        conn.close()
        
        if not quizzes:
            st.info("Chưa có quiz nào được tạo.")
        else:
            # Hiển thị danh sách quiz
            quiz_options = {f"{q['quiz_code']} - {q['title']} ({q['subject']})": q['id'] for q in quizzes}
            selected_quiz_label = st.selectbox("Chọn quiz để sửa:", list(quiz_options.keys()))
            selected_quiz_id = quiz_options[selected_quiz_label]
            
            # Lấy câu hỏi của quiz
            conn = sqlite3.connect('quiz_system.db')
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute("SELECT * FROM questions WHERE quiz_id = ? ORDER BY id", (selected_quiz_id,))
            questions = c.fetchall()
            conn.close()
            
            if not questions:
                st.warning("Quiz này chưa có câu hỏi.")
            else:
                st.subheader(f"✏️ Sửa câu hỏi cho quiz: {selected_quiz_label}")
                
                # Tạo form để sửa từng câu
                with st.form("edit_questions_form"):
                    updated_questions = []
                    for idx, q in enumerate(questions):
                        st.markdown(f"#### Câu {idx+1}")
                        # Sử dụng st.session_state để lưu giá trị hiện tại
                        question_text = st.text_area("Nội dung câu hỏi", value=q['question_text'], key=f"q_{q['id']}_text")
                        col1, col2 = st.columns(2)
                        with col1:
                            opt_a = st.text_input("Đáp án A", value=q['option_a'], key=f"q_{q['id']}_a")
                            opt_c = st.text_input("Đáp án C", value=q['option_c'], key=f"q_{q['id']}_c")
                        with col2:
                            opt_b = st.text_input("Đáp án B", value=q['option_b'], key=f"q_{q['id']}_b")
                            opt_d = st.text_input("Đáp án D", value=q['option_d'], key=f"q_{q['id']}_d")
                        correct = st.selectbox("Đáp án đúng", options=["A", "B", "C", "D"], index=["A","B","C","D"].index(q['correct_answer']), key=f"q_{q['id']}_correct")
                        explanation = st.text_area("Giải thích", value=q['explanation'], key=f"q_{q['id']}_expl")
                        st.markdown("---")
                        updated_questions.append({
                            'id': q['id'],
                            'question_text': question_text,
                            'option_a': opt_a,
                            'option_b': opt_b,
                            'option_c': opt_c,
                            'option_d': opt_d,
                            'correct_answer': correct,
                            'explanation': explanation
                        })
                    
                    submitted = st.form_submit_button("💾 LƯU THAY ĐỔI")
                    if submitted:
                        conn = sqlite3.connect('quiz_system.db')
                        c = conn.cursor()
                        for q_updated in updated_questions:
                            c.execute('''
                                UPDATE questions 
                                SET question_text=?, option_a=?, option_b=?, option_c=?, option_d=?, correct_answer=?, explanation=?
                                WHERE id=?
                            ''', (q_updated['question_text'], q_updated['option_a'], q_updated['option_b'], q_updated['option_c'], q_updated['option_d'], q_updated['correct_answer'], q_updated['explanation'], q_updated['id']))
                        conn.commit()
                        conn.close()
                        st.success("Đã cập nhật câu hỏi thành công!")
                        st.rerun()
    
    # ==================== XUẤT BÁO CÁO ====================
    elif menu == "📥 XUẤT BÁO CÁO":
        st.header("📥 XUẤT BÁO CÁO EXCEL")
        
        report_type = st.selectbox(
            "Chọn loại báo cáo:",
            [
                "📋 Toàn bộ kết quả",
                "🏫 Kết quả theo lớp",
                "📝 Kết quả theo Quiz", 
                "👨‍🎓 Kết quả học sinh"
            ]
        )
        
        if report_type == "🏫 Kết quả theo lớp":
            classes = get_distinct_classes()
            
            if classes:
                selected_classes = st.multiselect("Chọn lớp (có thể chọn nhiều):", classes)
                
                if selected_classes and st.button("📤 **XUẤT BÁO CÁO LỚP**", use_container_width=True):
                    conn = sqlite3.connect('quiz_system.db')
                    conn.row_factory = sqlite3.Row
                    c = conn.cursor()
                    
                    # Lấy dữ liệu
                    placeholders = ','.join(['?'] * len(selected_classes))
                    c.execute(f'''
                        SELECT * FROM results 
                        WHERE class_name IN ({placeholders})
                        ORDER BY class_name, student_name, submitted_at
                    ''', selected_classes)
                    
                    results = c.fetchall()
                    conn.close()
                    
                    if results:
                        # Chuẩn bị dữ liệu
                        data = []
                        for r in results:
                            data.append({
                                "Mã bài": r['id'],
                                "Họ tên": r['student_name'],
                                "Lớp": r['class_name'],
                                "Mã HS": r['student_id'] or "",
                                "Mã Quiz": r['quiz_code'],
                                "Điểm": r['score'],
                                "Tổng câu": r['total_questions'],
                                "Tỉ lệ %": r['percentage'],
                                "Xếp loại": r['grade'],
                                "Thời gian": r['submitted_at']
                            })
                        
                        df = pd.DataFrame(data)
                        
                        # Tạo Excel với nhiều sheet
                        excel_buffer = io.BytesIO()
                        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                            # Sheet chi tiết
                            df.to_excel(writer, index=False, sheet_name='Chi tiết')
                            
                            # Sheet thống kê
                            stats_data = []
                            for class_name in selected_classes:
                                class_df = df[df['Lớp'] == class_name]
                                if not class_df.empty:
                                    stats_data.append({
                                        "Lớp": class_name,
                                        "Số bài thi": len(class_df),
                                        "Số học sinh": class_df['Họ tên'].nunique(),
                                        "Điểm TB": f"{class_df['Tỉ lệ %'].mean():.1f}%",
                                        "Điểm cao nhất": f"{class_df['Tỉ lệ %'].max():.1f}%",
                                        "Điểm thấp nhất": f"{class_df['Tỉ lệ %'].min():.1f}%"
                                    })
                            
                            if stats_data:
                                pd.DataFrame(stats_data).to_excel(writer, index=False, sheet_name='Thống kê')
                        
                        excel_buffer.seek(0)
                        
                        st.success(f"✅ **ĐÃ XUẤT {len(results)} KẾT QUẢ CỦA {len(selected_classes)} LỚP**")
                        
                        # Nút download
                        st.download_button(
                            label="📥 **TẢI FILE EXCEL**",
                            data=excel_buffer,
                            file_name=f"bao_cao_lop_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
            else:
                st.info("📭 Chưa có dữ liệu lớp học")
        
        elif report_type == "📝 Kết quả theo Quiz":
            try:
                conn = sqlite3.connect('quiz_system.db')
                c = conn.cursor()
                c.execute("SELECT DISTINCT quiz_code FROM results ORDER BY quiz_code")
                quizzes = [row[0] for row in c.fetchall()]
                conn.close()
                
                if quizzes:
                    selected_quiz = st.selectbox("Chọn Quiz:", quizzes)
                    
                    if selected_quiz and st.button("📤 **XUẤT BÁO CÁO QUIZ**", use_container_width=True):
                        conn = sqlite3.connect('quiz_system.db')
                        conn.row_factory = sqlite3.Row
                        c = conn.cursor()
                        c.execute('SELECT * FROM results WHERE quiz_code = ? ORDER BY percentage DESC', (selected_quiz,))
                        results = c.fetchall()
                        conn.close()
                        
                        if results:
                            data = []
                            for r in results:
                                data.append({
                                    "Mã bài": r['id'],
                                    "Họ tên": r['student_name'],
                                    "Lớp": r['class_name'],
                                    "Mã HS": r['student_id'] or "",
                                    "Mã Quiz": r['quiz_code'],
                                    "Điểm": r['score'],
                                    "Tổng câu": r['total_questions'],
                                    "Tỉ lệ %": r['percentage'],
                                    "Xếp loại": r['grade'],
                                    "Thời gian": r['submitted_at']
                                })
                            
                            df = pd.DataFrame(data)
                            excel_buffer = io.BytesIO()
                            df.to_excel(excel_buffer, index=False, engine='openpyxl')
                            excel_buffer.seek(0)
                            
                            st.success(f"✅ **ĐÃ XUẤT {len(results)} KẾT QUẢ CỦA QUIZ {selected_quiz}**")
                            
                            st.download_button(
                                label="📥 **TẢI FILE EXCEL**",
                                data=excel_buffer,
                                file_name=f"bao_cao_quiz_{selected_quiz}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True
                            )
                else:
                    st.info("📭 Chưa có dữ liệu quiz")
                    
            except Exception as e:
                st.error(f"❌ Lỗi: {str(e)}")
        
        elif st.button("📤 **XUẤT TOÀN BỘ KẾT QUẢ**", use_container_width=True):
            conn = sqlite3.connect('quiz_system.db')
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute('SELECT * FROM results ORDER BY submitted_at DESC')
            results = c.fetchall()
            conn.close()
            
            if results:
                # Chuẩn bị dữ liệu
                data = []
                for r in results:
                    data.append({
                        "Mã bài": r['id'],
                        "Họ tên": r['student_name'],
                        "Lớp": r['class_name'],
                        "Mã HS": r['student_id'] or "",
                        "Mã Quiz": r['quiz_code'],
                        "Điểm": r['score'],
                        "Tổng câu": r['total_questions'],
                        "Tỉ lệ %": r['percentage'],
                        "Xếp loại": r['grade'],
                        "Thời gian": r['submitted_at']
                    })
                
                df = pd.DataFrame(data)
                excel_buffer = io.BytesIO()
                df.to_excel(excel_buffer, index=False, engine='openpyxl')
                excel_buffer.seek(0)
                
                st.success(f"✅ **ĐÃ XUẤT {len(results)} KẾT QUẢ**")
                
                st.download_button(
                    label="📥 **TẢI FILE EXCEL**",
                    data=excel_buffer,
                    file_name=f"toan_bo_ket_qua_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            else:
                st.info("📭 Chưa có dữ liệu để xuất")

if __name__ == "__main__":
    main()
